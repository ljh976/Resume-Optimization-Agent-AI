from io import BytesIO

import pytest
from docx import Document

from core.input_extract import InputExtractionError, extract_uploaded_text


def test_extract_txt_utf8():
    text = "Senior Data Engineer\nPython, SQL, and Airflow required."
    assert extract_uploaded_text("role.txt", text.encode("utf-8")) == text


def test_extract_docx_includes_paragraphs_and_tables():
    document = Document()
    document.add_paragraph("Backend Engineer")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Required"
    table.cell(0, 1).text = "Python and AWS"
    payload = BytesIO()
    document.save(payload)

    result = extract_uploaded_text("role.docx", payload.getvalue())

    assert "Backend Engineer" in result
    assert "Required | Python and AWS" in result


def test_extract_pdf():
    pytest.importorskip("pypdf")
    reportlab_canvas = pytest.importorskip("reportlab.pdfgen.canvas")
    payload = BytesIO()
    canvas = reportlab_canvas.Canvas(payload)
    canvas.drawString(72, 720, "Platform Engineer - Kubernetes and Terraform")
    canvas.save()

    result = extract_uploaded_text("role.pdf", payload.getvalue())

    assert "Platform Engineer" in result
    assert "Kubernetes" in result


def test_rejects_unsupported_or_empty_files():
    with pytest.raises(InputExtractionError, match="Unsupported"):
        extract_uploaded_text("role.rtf", b"job description")
    with pytest.raises(InputExtractionError, match="empty"):
        extract_uploaded_text("role.txt", b"")
