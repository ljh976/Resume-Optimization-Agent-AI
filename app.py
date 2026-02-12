import streamlit as st, tempfile, os, time, math
from core.agent import generate, evaluate, rewrite, trimmer, _heuristic_ats_score
from core.header_extract import extract_header_info, build_header_lines
from core.prescreen import prescreen_resume
from core.structure import parse_resume, split_experience, merge_skills_a1
from core.render import render_docx
# relevance scorer
# per-bullet LLM scorer removed to reduce LLM calls; use heuristic scoring if needed
# static career pages list
from core.career_pages import ALL_COMPANIES, CATEGORIZED
# recommend_jobs removed

# PDF conversion removed; we estimate pages by word count only

st.set_page_config(page_title="Resume Optimization Agent", layout="wide")
st.title("Resume Optimization Agent")

# Persisted state helpers: save/load inputs + last result to a local JSON file
import json

_STATE_PATH = os.path.join(os.getcwd(), '.resume_agent_state.json')

def _inputs_hash(jd_value: str, master_value: str) -> str:
    import hashlib
    payload = (jd_value or "") + "\n---\n" + (master_value or "")
    return hashlib.sha256(payload.encode('utf-8')).hexdigest()

def _load_persisted_state():
    try:
        if st.session_state.get('_state_loaded'):
            return
        if os.path.exists(_STATE_PATH):
            with open(_STATE_PATH, 'r', encoding='utf-8') as f:
                data = json.load(f)
            # restore inputs
            if isinstance(data, dict):
                if 'jd_input' in data and data.get('jd_input') is not None:
                    st.session_state.setdefault('jd_input', data.get('jd_input'))
                if 'master_input' in data and data.get('master_input') is not None:
                    st.session_state.setdefault('master_input', data.get('master_input'))
                # restore result if present and input hash matches
                if 'result' in data and isinstance(data.get('result'), dict):
                    r = data.get('result')
                    resume_text = r.get('resume')
                    fb = r.get('fb')
                    saved_hash = r.get('input_hash')
                    current_hash = _inputs_hash(
                        st.session_state.get('jd_input') or "",
                        st.session_state.get('master_input') or ""
                    )
                    if saved_hash and saved_hash == current_hash and resume_text is not None and fb is not None:
                        st.session_state.setdefault('result', (resume_text, fb))
                if 'suggested_skills' in data and data.get('suggested_skills') is not None:
                    st.session_state.setdefault('suggested_skills', data.get('suggested_skills') or [])
                if 'generated_bullet_scores' in data and data.get('generated_bullet_scores') is not None:
                    st.session_state.setdefault('generated_bullet_scores', data.get('generated_bullet_scores') or [])
        st.session_state['_state_loaded'] = True
    except Exception:
        pass




def _save_persisted_state():
    try:
        to_save = {}
        to_save['jd_input'] = st.session_state.get('jd_input')
        to_save['master_input'] = st.session_state.get('master_input')
        res = st.session_state.get('result')
        if res and isinstance(res, (list, tuple)) and len(res) == 2:
            resume_text, fb = res
            # Always save resume_text; try to include fb but fall back to None if not serializable
            try:
                # attempt to serialize feedback; if it fails, store None for fb
                _ = json.dumps(fb)
                fb_serializable = fb
            except Exception:
                try:
                    fb_serializable = str(fb)
                except Exception:
                    fb_serializable = None
            to_save['result'] = {
                'resume': resume_text,
                'fb': fb_serializable,
                'input_hash': _inputs_hash(
                    st.session_state.get('jd_input') or "",
                    st.session_state.get('master_input') or ""
                )
            }
        else:
            to_save['result'] = None
        to_save['suggested_skills'] = st.session_state.get('suggested_skills')
        to_save['generated_bullet_scores'] = st.session_state.get('generated_bullet_scores')
        with open(_STATE_PATH, 'w', encoding='utf-8') as f:
            json.dump(to_save, f, ensure_ascii=False)
    except Exception:
        pass

def _clear_cached_result():
    try:
        st.session_state.pop('result', None)
        st.session_state.pop('suggested_skills', None)
        st.session_state.pop('generated_bullet_scores', None)
    except Exception:
        pass
    try:
        _save_persisted_state()
    except Exception:
        pass

def _on_input_change():
    _clear_cached_result()
    _save_persisted_state()

def _normalize_master_with_header(master_text: str, header_info: dict) -> str:
    header_lines = build_header_lines(header_info or {})
    if len(header_lines) <= 1:
        return master_text
    if (master_text or "").lstrip().upper().startswith("HEADER"):
        return master_text
    header_block = "\n".join(header_lines)
    return header_block + "\n" + master_text

def _clear_inputs():
    st.session_state['pending_input_clear'] = True
    st.session_state['clear_inputs_notice'] = True

# Load persisted values early so widgets pick them up
_load_persisted_state()

if st.session_state.pop('pending_input_clear', False):
    st.session_state['jd_input'] = ''
    st.session_state['master_input'] = ''
    try:
        _save_persisted_state()
    except Exception:
        pass

# Default test_mode to True unless persisted otherwise
st.session_state.setdefault('test_mode', True)

# Safety: when True, prevent any programmatic trimming/forced removals so
# the agent never deletes content outside the LLM's control.
DISABLE_PROGRAMMATIC_TRIMS = True

def _get_docx_bytes(resume_text):
    import tempfile
    try:
        secs = parse_resume(resume_text)
        roles_local = split_experience(secs.get('EXPERIENCE', []))
        skills_local = merge_skills_a1(secs.get('SKILLS', []))
    except Exception:
        secs = {'EXPERIENCE': [], 'SKILLS': []}
        roles_local = []
        skills_local = []

    try:
        # If parsing produced no useful structured sections (e.g., the model only emitted SKILLS
        # but no EXPERIENCE/SUMMARY/EDUCATION), fall back to rendering the raw resume text to
        # preserve all content rather than lose body paragraphs.
        useful_sections = 0
        for sname in ('SUMMARY', 'EXPERIENCE', 'EDUCATION'):
            if secs.get(sname):
                useful_sections += 1
        if useful_sections == 0:
            from docx import Document
            doc = Document()
            p = doc.add_paragraph()
            p.add_run(resume_text or "")
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tf:
                tmp_path = tf.name
            doc.save(tmp_path)
            with open(tmp_path, 'rb') as f:
                data = f.read()
            try:
                os.remove(tmp_path)
            except Exception:
                pass
            return data

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tf:
            tmp_path = tf.name
        render_docx(tmp_path, secs, roles_local, skills_local)
        with open(tmp_path, 'rb') as f:
            data = f.read()
        try:
            os.remove(tmp_path)
        except Exception:
            pass
        return data
    except Exception:
        return None


def _estimate_resume_lines(resume_text: str) -> int:
    """Estimate Word line usage for a structured resume.

    Excludes name + contact line (HEADER content). Used only as a heuristic to
    enforce a one-page constraint via a 40-line budget.
    """
    try:
        secs = parse_resume(resume_text or "")
        roles_local = split_experience(secs.get("EXPERIENCE", []) or secs.get("PROFESSIONAL EXPERIENCE", []) or [])
        merged_skills_local = merge_skills_a1(secs.get("SKILLS", []) or [])
    except Exception:
        return len([l for l in (resume_text or "").splitlines() if l.strip()])

    def _wrap_penalty(text: str, threshold: int = 125) -> int:
        t = (text or "").strip()
        if not t:
            return 0
        if len(t) <= threshold:
            return 0
        if len(t) <= threshold * 2:
            return 1
        return 2

    lines = 0

    # SUMMARY
    if secs.get("SUMMARY") or secs.get("PROFESSIONAL SUMMARY"):
        lines += 1
        lines += len(secs.get("SUMMARY", []) or secs.get("PROFESSIONAL SUMMARY", []) or [])

    # EXPERIENCE
    if secs.get("EXPERIENCE") or secs.get("PROFESSIONAL EXPERIENCE"):
        lines += 1
        if roles_local:
            for r in roles_local:
                lines += 1
                for b in (r.get("bullets") or []):
                    lines += 1 + _wrap_penalty(b)
        else:
            exp_lines = secs.get("EXPERIENCE", []) or secs.get("PROFESSIONAL EXPERIENCE", []) or []
            for l in exp_lines:
                s = (l or "").strip()
                if not s:
                    continue
                if s.startswith("-") or s.startswith("•"):
                    lines += 1 + _wrap_penalty(s[1:].strip())
                else:
                    lines += 1

    # SKILLS
    if secs.get("SKILLS"):
        lines += 1
        lines += len([l for l in merged_skills_local if isinstance(l, str) and l.strip()])

    # EDUCATION
    if secs.get("EDUCATION"):
        lines += 1
        lines += len([l for l in (secs.get("EDUCATION", []) or []) if isinstance(l, str) and l.strip()])

    # Baseline overhead: Word paragraph spacing + underline rendering + minor wrap variance.
    # Empirically this tends to add ~2 lines beyond our plain-text line estimate.
    baseline_overhead = 2
    return int(lines + baseline_overhead)


def _trim_resume_to_max_lines(resume_text: str, max_lines: int = 40) -> str:
    """Deterministically trim a structured resume to fit within `max_lines`.

    Priority is to drop EXPERIENCE bullets first, then SKILLS, then SUMMARY,
    keeping sections present whenever possible.
    """
    try:
        secs = parse_resume(resume_text or "")
        header_lines = secs.get('HEADER', []) or []
        summary_lines = (secs.get('SUMMARY', []) or secs.get('PROFESSIONAL SUMMARY', []) or []).copy()
        exp_lines = (secs.get('EXPERIENCE', []) or secs.get('PROFESSIONAL EXPERIENCE', []) or []).copy()
        skills_lines = merge_skills_a1(secs.get('SKILLS', []) or [])
        edu_lines = (secs.get('EDUCATION', []) or []).copy()
    except Exception:
        kept = []
        for l in (resume_text or "").splitlines():
            if l.strip():
                kept.append(l)
            if len(kept) >= max_lines:
                break
        return "\n".join(kept).strip()

    roles_local = split_experience(exp_lines)

    def _build_from_roles(roles_list, summary_subset, skills_subset, edu_subset):
        parts = []
        parts.append('HEADER')
        parts.extend(header_lines)
        parts.append('SUMMARY')
        parts.extend(summary_subset)
        parts.append('EXPERIENCE')
        if roles_list:
            for r in roles_list:
                hdr = ""
                if r.get('company') and r.get('meta'):
                    hdr = f"{r.get('company')} | {r.get('meta')}"
                elif r.get('company'):
                    hdr = str(r.get('company'))
                elif r.get('meta'):
                    hdr = str(r.get('meta'))
                if hdr:
                    parts.append(hdr)
                for b in (r.get('bullets') or []):
                    parts.append('- ' + str(b).strip())
        else:
            parts.extend(exp_lines)
        parts.append('SKILLS')
        parts.extend([l for l in (skills_subset or []) if isinstance(l, str) and l.strip()])
        parts.append('EDUCATION')
        parts.extend([l for l in (edu_subset or []) if isinstance(l, str) and l.strip()])
        return "\n".join([p for p in parts if p is not None and str(p).strip() != ""]).strip()

    safety = 0
    while _estimate_resume_lines(_build_from_roles(roles_local, summary_lines, skills_lines, edu_lines)) > max_lines and safety < 500:
        safety += 1
        changed = False

        # Drop last EXPERIENCE bullet (keep at least 1 bullet per role when possible)
        for ri in range(len(roles_local) - 1, -1, -1):
            bl = roles_local[ri].get('bullets') or []
            if len(bl) > 1:
                bl.pop(-1)
                roles_local[ri]['bullets'] = bl
                changed = True
                break
        if changed:
            continue

        # Drop a SKILLS line
        if skills_lines and len(skills_lines) > 1:
            skills_lines = skills_lines[:-1]
            changed = True
        if changed:
            continue

        # Drop SUMMARY line but keep at least 3 if present
        if summary_lines and len(summary_lines) > 3:
            summary_lines = summary_lines[:-1]
            changed = True
        if changed:
            continue

        # Drop a whole role only if it has no bullets left (avoid collapsing experience)
        for ri in range(len(roles_local) - 1, -1, -1):
            bl = roles_local[ri].get('bullets') or []
            if len(bl) <= 1 and len(roles_local) > 1:
                roles_local.pop(ri)
                changed = True
                break
        if changed:
            continue

        # As last resort, drop EDUCATION lines but keep at least 1
        if edu_lines and len(edu_lines) > 1:
            edu_lines = edu_lines[:-1]
            changed = True
        if not changed:
            break

    return _build_from_roles(roles_local, summary_lines, skills_lines, edu_lines)


def _truncate_structured_resume(resume_text: str, cap: int = 12000) -> str:
    """Emergency-only cap for extreme outputs; keeps structured headers if possible."""
    if not resume_text:
        return ""
    if len(resume_text) <= cap:
        return resume_text.rstrip()

    try:
        secs = parse_resume(resume_text)
    except Exception:
        return resume_text[:cap].rstrip()

    header_lines = secs.get('HEADER', []) or []
    summary_lines = (secs.get('SUMMARY', []) or secs.get('PROFESSIONAL SUMMARY', []) or []).copy()
    exp_lines = (secs.get('EXPERIENCE', []) or secs.get('PROFESSIONAL EXPERIENCE', []) or []).copy()
    skills_lines = (secs.get('SKILLS', []) or []).copy()
    edu_lines = (secs.get('EDUCATION', []) or []).copy()

    def _build() -> str:
        parts = []
        parts.append('HEADER')
        parts.extend(header_lines)
        parts.append('SUMMARY')
        parts.extend(summary_lines)
        parts.append('EXPERIENCE')
        parts.extend(exp_lines)
        parts.append('SKILLS')
        parts.extend(skills_lines)
        parts.append('EDUCATION')
        parts.extend(edu_lines)
        return "\n".join([p for p in parts if p is not None and str(p).strip() != ""]).strip()

    def _is_bullet(line: str) -> bool:
        s = (line or "").lstrip()
        return s.startswith('-') or s.startswith('•')

    min_summary = 1 if summary_lines else 0
    min_exp = 1 if exp_lines else 0
    min_skills = 1 if skills_lines else 0
    min_edu = 1 if edu_lines else 0

    def _shrink_skills_items_once(*, max_len_hint: int = 140) -> bool:
        """Prefer trimming SKILLS items (fine-grained) over dropping whole bullets.

        This reduces the chance we overshoot far below `cap` by removing an entire
        long EXPERIENCE bullet or role header.
        """
        nonlocal skills_lines
        if not skills_lines:
            return False

        # Pick the longest skills line as the best candidate to shrink.
        try:
            idx = max(range(len(skills_lines)), key=lambda i: len((skills_lines[i] or "").strip()))
        except Exception:
            return False

        line = (skills_lines[idx] or "").strip()
        if not line:
            return False
        if len(line) <= max_len_hint:
            return False

        delim = None
        if ":" in line:
            delim = ":"
        elif "：" in line:
            delim = "："
        if not delim:
            return False

        cat, rest = line.split(delim, 1)
        cat = cat.strip()
        items = [i.strip() for i in (rest or "").split(",") if i.strip()]
        if not cat or len(items) <= 1:
            # If we cannot shrink this line further, allow dropping a whole SKILLS line later.
            return False

        # Drop trailing items until the line is shorter (one item per call).
        items.pop(-1)
        skills_lines[idx] = f"{cat}: " + ", ".join(items)
        return True

    def _drop_last(arr, *, min_keep: int) -> bool:
        if len(arr) <= min_keep:
            return False
        arr.pop(-1)
        return True

    def _drop_last_matching(arr, pred, *, min_keep: int) -> bool:
        if len(arr) <= min_keep:
            return False
        for i in range(len(arr) - 1, -1, -1):
            if pred(arr[i]):
                arr.pop(i)
                return True
        return False

    safety = 0
    while len(_build()) > cap and safety < 800:
        safety += 1
        changed = False

        # First: attempt fine-grained SKILLS shrinking to avoid large content drops.
        changed = _shrink_skills_items_once() or changed
        if changed:
            continue

        # Prefer dropping EXPERIENCE bullets first
        changed = _drop_last_matching(exp_lines, _is_bullet, min_keep=min_exp) or changed
        if changed:
            continue

        # Then non-bullet EXPERIENCE lines
        changed = _drop_last_matching(exp_lines, lambda x: not _is_bullet(x), min_keep=min_exp) or changed
        if changed:
            continue

        # Then SKILLS
        changed = _drop_last(skills_lines, min_keep=min_skills) or changed
        if changed:
            continue

        # Then SUMMARY
        changed = _drop_last(summary_lines, min_keep=min_summary) or changed
        if changed:
            continue

        # Then EDUCATION
        changed = _drop_last(edu_lines, min_keep=min_edu) or changed
        if not changed:
            break

    built = _build()
    if len(built) > cap:
        return built[:cap].rstrip()
    return built


# Helper: compute total years from parsed roles (derive earliest start and latest end)
def _compute_total_years_from_roles(roles):
    import re
    import datetime

    starts = []
    ends = []
    now_year = datetime.datetime.now().year
    year_re = re.compile(r"(\d{4})")
    range_re = re.compile(r"(\d{4})\s*[–-]\s*(\d{4}|Present|Now|present|now)")

    for r in roles:
        meta = r.get("meta", "")
        m = range_re.search(meta)
        if m:
            try:
                s = int(m.group(1))
                egrp = m.group(2)
                if egrp.isdigit():
                    e = int(egrp)
                else:
                    e = now_year
                starts.append(s)
                ends.append(e)
                continue
            except Exception:
                pass

        ys = year_re.findall(meta)
        if ys:
            try:
                s = int(ys[0])
                starts.append(s)
                if len(ys) > 1:
                    ends.append(int(ys[-1]))
                else:
                    ends.append(now_year)
            except Exception:
                continue

    if not starts or not ends:
        return None

    start = min(starts)
    end = max(ends)
    total = end - start
    if total < 0:
        return None
    return total


def _extract_summary_year_claim(resume_text):
    import re
    # look for patterns like '10+ years', '10 years', 'over 10 years'
    m = re.search(r"(\d{1,2})\s*\+?\s*(?:\+\s*)?(?:years|yrs|year)", resume_text, re.IGNORECASE)
    if m:
        try:
            return int(m.group(1))
        except Exception:
            return None
    return None
    return None


def _safe_rerun():
    """Attempt to rerun the Streamlit script in a version-compatible way.
    Tries `st.experimental_rerun()` first; if unavailable, raises the
    Streamlit RerunException from known import paths. Falls back to toggling
    a session_state flag if rerun cannot be forced.
    """
    if hasattr(st, 'experimental_rerun'):
        st.experimental_rerun()
        return

    # Try known exception class locations to force a rerun when experimental_rerun
    # is unavailable (older Streamlit versions).
    RerunException = None
    rerun_ctx = None
    for path in [
        'streamlit.runtime.scriptrunner.script_runner',
        'streamlit.script_runner.script_runner'
    ]:
        try:
            module = __import__(path, fromlist=['RerunException'])
            RerunException = getattr(module, 'RerunException', None)
            if RerunException:
                rerun_ctx = getattr(module, 'RerunData', None)
                break
        except Exception:
            continue

    if RerunException:
        if rerun_ctx:
            raise RerunException(rerun_ctx())
        raise RerunException(None)

    # Last resort: toggle a flag so next interaction triggers a rerun.
    st.session_state['_need_rerun'] = not st.session_state.get('_need_rerun', False)

# Create a two-column top layout: main inputs on the left, persistent
# controls on the right (simulates a right-side sidebar)
main_col, sidebar_col = st.columns([3, 1])

# Ensure running flag exists so widgets disabled logic is deterministic
st.session_state.setdefault('running', False)
st.session_state.setdefault('run_requested', False)

with main_col:
    st.subheader("Inputs")
    with st.expander("Job Description", expanded=True):
        jd = st.text_area("Job Description", height=220, key="jd_input", on_change=_save_persisted_state)
    with st.expander("Master Resume", expanded=True):
        master = st.text_area("Master Resume", height=320, key="master_input", on_change=_save_persisted_state)

    # Place Run Agent under the Master Resume field as requested
    def _on_run_click():
        try:
            st.session_state['run_requested'] = True
        except Exception:
            st.session_state.__setitem__('run_requested', True)
        # No rerun here; Streamlit will rerun after callback naturally.

    run_clicked = st.button("Run Agent", disabled=st.session_state.get('running', False), on_click=_on_run_click)
    # placeholder for status messages shown under the Run button (main column)
    status_placeholder = main_col.empty()
    # persistent progress placeholder so reruns still show progress
    main_progress_placeholder = main_col.empty()

    # if a run is already in progress, re-render progress UI from session state
    if st.session_state.get('running'):
        pct = st.session_state.get('progress_pct', 0)
        try:
            main_progress_placeholder.progress(min(100, max(0, int(pct))))
        except Exception:
            pass

        # (Career Pages removed from the running-only area; rendered persistently at bottom)
        msg = st.session_state.get('progress_msg')
        if msg:
            try:
                status_placeholder.write(f"{msg} — elapsed: {int(time.time() - st.session_state.get('run_start_ts', time.time()))}s")
            except Exception:
                status_placeholder.write(msg)

with sidebar_col:
    st.header("Controls")
    target = st.slider("Target ATS Score", 60, 95, 80, disabled=st.session_state.get('running', False))
    max_iters = st.slider("Max Iterations", 1, 6, 1, disabled=st.session_state.get('running', False))
    # One-page trimming removed per user request. Keep UI minimal.
    auto_trim = False
    show_trim_logs = False
    test_mode = st.checkbox("Test mode (use cheapest LLM model)", value=True, help="Run using the fastest/cheapest model for LLM calls (still performs API calls).", disabled=st.session_state.get('running', False))
    prescreen_enabled = st.checkbox(
        "Pre-screening enabled",
        value=True,
        help="Stop early if skill match is too low.",
        disabled=st.session_state.get('running', False),
        key="prescreen_enabled"
    )
    # job-recommendation controls removed

    # Company recommendations are derived automatically from the resume.

    if st.button("Clear Cached Input", key="clear_inputs_btn", disabled=st.session_state.get('running', False)):
        _clear_inputs()
    if st.session_state.pop('clear_inputs_notice', False):
        st.success("Cached inputs cleared.")
    if st.button("Clear Cached Result", key="clear_cache_btn", disabled=st.session_state.get('running', False)):
        _clear_cached_result()
        st.success("Cached result cleared.")

    # Career pages listing moved below (rendered after Feedback section)
    # (placeholder here so controls stay grouped at top)

# Initialize session state for recommendations if missing
# job_recs removed; recommendations feature disabled

    # Start run if button clicked or a run was requested via on_click
run_requested = st.session_state.pop('run_requested', False)
if run_clicked or run_requested:
    # validate inputs first
    if not jd or not jd.strip() or not master or not master.strip():
        status_placeholder.error("Please provide both Job Description and Master Resume before running the agent.")
    else:
        # show status messages under the Run Agent button
        status_placeholder.info("Running optimization agent...")
        # prevent double runs and reset summary-fix attempt for this run
        st.session_state['running'] = True
        st.session_state['summary_fix_attempted'] = False
        st.session_state['prescreen_result'] = None
        st.session_state['loop_scores'] = []
        _clear_cached_result()
        # progress UI: small progress bar and elapsed timer (persist via session_state)
        st.session_state['run_start_ts'] = time.time()
        st.session_state['progress_pct'] = 0
        st.session_state['progress_msg'] = None
        try:
            main_progress_placeholder.progress(0)
        except Exception:
            pass

        def _update_progress(pct, msg=None):
            try:
                main_progress_placeholder.progress(min(100, max(0, int(pct))))
            except Exception:
                pass
            st.session_state['progress_pct'] = int(pct) if pct is not None else st.session_state.get('progress_pct', 0)
            st.session_state['progress_msg'] = msg
            if msg:
                try:
                    elapsed = int(time.time() - st.session_state.get('run_start_ts', time.time()))
                    status_placeholder.write(f"{msg} — elapsed: {elapsed}s")
                except Exception:
                    status_placeholder.write(msg)

        def _report_error(msg: str, exc: Exception = None):
            try:
                status_placeholder.error(msg)
            except Exception:
                pass
            if exc is not None:
                try:
                    st.exception(exc)
                except Exception:
                    pass

        # Safe caller for potentially long LLM operations to avoid indefinite blocking
        def _safe_call(fn, *a, timeout=60, **kw):
            """Run `fn(*a, **kw)` in a background thread and return result or raise on timeout.
            Important: do not block waiting for the worker to finish after a timeout (avoid join hang).
            """
            from concurrent.futures import ThreadPoolExecutor, TimeoutError as FuturesTimeoutError
            ex = ThreadPoolExecutor(max_workers=1)
            fut = ex.submit(fn, *a, **kw)
            try:
                return fut.result(timeout=timeout)
            except FuturesTimeoutError:
                # attempt to cancel and shutdown without waiting to avoid blocking the main thread
                try:
                    fut.cancel()
                except Exception:
                    pass
                try:
                    ex.shutdown(wait=False)
                except Exception:
                    pass
                raise TimeoutError(f"LLM call to {getattr(fn, '__name__', str(fn))} timed out after {timeout}s")
            finally:
                # safe shutdown if still running; do not block
                try:
                    ex.shutdown(wait=False)
                except Exception:
                    pass

        _update_progress(5, "Starting")
        # Fail fast if API key missing
        if not os.getenv('OPENAI_API_KEY'):
            status_placeholder.error("OPENAI_API_KEY is not set. Set the env variable and restart the app.")
            st.session_state['running'] = False
            _safe_rerun()
            st.stop()

        # Quick probe: check OpenAI API reachable and key valid (fast fail)
        def _probe_openai(timeout=5):
            import requests, os
            k = os.getenv('OPENAI_API_KEY')
            if not k:
                return False, "OPENAI_API_KEY not set"
            try:
                r = requests.get('https://api.openai.com/v1/models', headers={'Authorization': f'Bearer {k}'}, timeout=timeout)
                if r.status_code == 200:
                    return True, None
                # return API error message if available
                try:
                    data = r.json()
                    msg = data.get('error', {}).get('message') if isinstance(data, dict) else None
                except Exception:
                    msg = r.text
                return False, f"OpenAI probe failed: {r.status_code} - {msg}"
            except requests.exceptions.RequestException as e:
                return False, str(e)

        ok, probe_msg = _probe_openai(timeout=5)
        offline_mode = False
        if not ok:
            # Instead of stopping, fall back to offline heuristic mode so the app remains usable.
            status_placeholder.warning(f"OpenAI API probe failed: {probe_msg} — continuing in offline mode (no LLM).")
            offline_mode = True
        st.session_state["offline_mode"] = offline_mode

        # Extract header info first to stabilize parsing for different input formats
        try:
            header_info = extract_header_info(master, use_llm=not offline_mode)
        except Exception:
            header_info = extract_header_info(master, use_llm=False)
        normalized_master = _normalize_master_with_header(master, header_info)

        # Pre-screening (skill match) before any optimization loop
        if prescreen_enabled:
            try:
                _update_progress(7, "Pre-screening")
                prescreen = prescreen_resume(jd, master, use_llm=not offline_mode)
            except Exception:
                prescreen = prescreen_resume(jd, master, use_llm=False)
            st.session_state['prescreen_result'] = prescreen
            if not prescreen.get("viable"):
                status_placeholder.warning(
                    f"Pre-screening stopped (skill match {prescreen.get('skill_match_pct', 0)}%). "
                    "Matching elements are too limited to optimize reliably."
                )
                tips = prescreen.get("tips") or []
                for t in tips[:2]:
                    status_placeholder.info(t)
                st.session_state['running'] = False
                _safe_rerun()
                st.stop()
            else:
                try:
                    status_placeholder.write(
                        f"Pre-screening passed (skill match {prescreen.get('skill_match_pct', 0)}%)."
                    )
                except Exception:
                    pass
        else:
            st.session_state['prescreen_result'] = None

        # If test mode is enabled, override the model to a cheap/fast model for all LLM calls
        model_overridden = False
        orig_model = os.environ.get('OPENAI_MODEL')
        if st.session_state.get('test_mode') or test_mode:
            os.environ['OPENAI_MODEL'] = 'gpt-3.5-turbo'
            model_overridden = True

        try:
            if offline_mode:
                # Build a minimal generator output so downstream parsing still works
                import json
                status_placeholder.warning("Offline mode: skipping LLM generate(); using MASTER as draft.")
                _update_progress(10, "Offline draft")
                raw = json.dumps({"suggested_skills": [], "bullet_scores": []}) + "\n" + normalized_master
            else:
                status_placeholder.info("Calling LLM generate() (timeout 60s)...")
                _update_progress(10, "Calling generator")
                raw = _safe_call(generate, jd, normalized_master, timeout=60)
                status_placeholder.success("Generator returned.")
                _update_progress(15, "Initial draft generated")
        except TimeoutError as e:
            _report_error("LLM generate timed out", e)
            st.session_state['running'] = False
            _safe_rerun()
            st.stop()
        except Exception as e:
            _report_error("Generator failed", e)
            st.session_state['running'] = False
            _safe_rerun()
            st.stop()
        finally:
            # restore model env var if we overrode it
            if model_overridden:
                if orig_model is not None:
                    os.environ['OPENAI_MODEL'] = orig_model
                else:
                    os.environ.pop('OPENAI_MODEL', None)
        # Attempt to extract JSON blob {"suggested_skills": [...]} appended after resume
        def _extract_json_blob(s: str):
            import json
            for start_ch in ('{', '['):
                start_idx = s.find(start_ch)
                if start_idx == -1:
                    continue
                stack = []
                pairs = {'{': '}', '[': ']'}
                open_ch = start_ch
                close_ch = pairs[open_ch]
                for k in range(start_idx, len(s)):
                    if s[k] == open_ch:
                        stack.append(open_ch)
                    elif s[k] == close_ch:
                        stack.pop()
                        if not stack:
                            end_idx = k + 1
                            candidate = s[start_idx:end_idx]
                            before = s[:start_idx]
                            try:
                                parsed = json.loads(candidate)
                                return parsed, before
                            except Exception:
                                break
            return None, s

        parsed_json, resume = _extract_json_blob(raw)
        suggested_skills_for_ui = []
        generated_bullet_scores = []
        if parsed_json and isinstance(parsed_json, dict):
            suggested_skills_for_ui = parsed_json.get('suggested_skills', []) or []
            generated_bullet_scores = parsed_json.get('bullet_scores', []) or []
        st.session_state['suggested_skills'] = suggested_skills_for_ui
        st.session_state['generated_bullet_scores'] = generated_bullet_scores
        _update_progress(15, "Initial draft generated")

        # After generation, validate SUMMARY years against MASTER-derived years
        try:
            sections_tmp = parse_resume(resume)
            roles_tmp = split_experience(sections_tmp.get("EXPERIENCE", []))
            computed_years = _compute_total_years_from_roles(roles_tmp)
            summary_claim = _extract_summary_year_claim(resume)
            if summary_claim and computed_years is not None and summary_claim > computed_years and not st.session_state.get('summary_fix_attempted'):
                # request a rewrite to fix summary (don't loop infinitely)
                fb_fix = {"fix_summary": True, "computed_years": computed_years}
                fb_fix['preserve_style'] = True
                fb_fix['prefer_shorten'] = True
                try:
                    raw_fix = _safe_call(rewrite, jd, normalized_master, resume, fb_fix, timeout=60)
                except Exception:
                    raw_fix = rewrite(jd, normalized_master, resume, fb_fix)
                st.session_state['summary_fix_attempted'] = True
                # try to extract JSON then resume text (reuse existing parser)
                def _extract_json_blob_local2(s: str):
                    import json
                    for start_ch in ('{', '['):
                        start_idx = s.find(start_ch)
                        if start_idx == -1:
                            continue
                        stack = []
                        pairs = {'{': '}', '[': ']'}
                        open_ch = start_ch
                        close_ch = pairs[open_ch]
                        for k in range(start_idx, len(s)):
                            if s[k] == open_ch:
                                stack.append(open_ch)
                            elif s[k] == close_ch:
                                stack.pop()
                                if not stack:
                                    end_idx = k + 1
                                    candidate = s[start_idx:end_idx]
                                    rest = s[end_idx:]
                                    try:
                                        parsed = json.loads(candidate)
                                        return parsed, rest
                                    except Exception:
                                        break
                    return None, s

                parsed_json_fix, remaining_fix = _extract_json_blob_local2(raw_fix)
                if parsed_json_fix is not None:
                    resume = remaining_fix.strip()
                else:
                    resume = raw_fix

        except Exception as e:
            _report_error("Summary validation failed", e)

        
        # Single evaluation pass — base score
        try:
            fb = _safe_call(evaluate, jd, resume, None, timeout=60)
        except Exception:
            fb = evaluate(jd, resume)
        _update_progress(30, f"ATS score: {fb.get('ats_score')}")

        # Helper to extract a JSON blob and return (parsed, remaining_text)
        def _extract_json_blob_and_rest(s: str):
            import json
            for start_ch in ('{', '['):
                start_idx = s.find(start_ch)
                if start_idx == -1:
                    continue
                stack = []
                pairs = {'{': '}', '[': ']'}
                open_ch = start_ch
                close_ch = pairs[open_ch]
                for k in range(start_idx, len(s)):
                    if s[k] == open_ch:
                        stack.append(open_ch)
                    elif s[k] == close_ch:
                        stack.pop()
                        if not stack:
                            end_idx = k + 1
                            candidate = s[start_idx:end_idx]
                            rest = s[end_idx:]
                            try:
                                parsed = json.loads(candidate)
                                return parsed, rest
                            except Exception:
                                break
            return None, s

        # Score-improvement loop: call `rewrite()` up to `max_iters` times to reach target
        try:
            for it in range(max_iters):
                current_score = int(fb.get('ats_score') or 0)
                if current_score >= int(target):
                    break
                prev_resume = resume
                prev_fb = dict(fb or {})
                prev_len = len(resume or "")
                prev_score = current_score
                _update_progress(30 + int((it + 1) / max_iters * 30), f"Improving score (iter {it+1}/{max_iters})")
                fb_update = dict(fb or {})
                fb_update['improve_score'] = True
                fb_update['target_score'] = int(target)
                fb_update.setdefault('preserve_style', True)
                fb_update.setdefault('prefer_shorten', False)

                try:
                    raw_rewrite = _safe_call(rewrite, jd, normalized_master, resume, fb_update, timeout=60)
                except Exception:
                    raw_rewrite = rewrite(jd, normalized_master, resume, fb_update)
                parsed_rewrite, remaining_rewrite = _extract_json_blob_and_rest(raw_rewrite)
                if parsed_rewrite is not None:
                    # attach any returned change log into feedback for UI
                    if isinstance(parsed_rewrite, dict):
                        fb.update(parsed_rewrite)
                    resume = remaining_rewrite.strip()
                else:
                    resume = raw_rewrite

                # re-evaluate after this improvement rewrite
                try:
                    fb = _safe_call(evaluate, jd, resume, None, timeout=40)
                except Exception:
                    try:
                        fb = evaluate(jd, resume)
                    except Exception:
                        pass

                # Guardrail: prevent iterations that aggressively shorten content without score gain.
                try:
                    new_score = int((fb or {}).get('ats_score') or 0)
                    new_len = len(resume or "")
                    if (prev_len - new_len) >= 250 and new_score <= prev_score:
                        resume = prev_resume
                        fb = prev_fb
                except Exception:
                    pass
                _update_progress(30 + int((it + 1) / max_iters * 30), f"ATS score: {fb.get('ats_score')}")
        except Exception as e:
            _report_error("Improvement loop failed", e)

        # Programmatic trimming / one-page reduction has been removed by user request.

        # If the resume is too short to visually fill one page in DOCX, do a single
        # repopulation pass that is allowed to restore factual bullets from MASTER.
        # This avoids "underfilled" pages without changing formatting.
        try:
            # If output is getting too short, run one repopulation pass to restore factual content.
            fill_min_chars = 2450
            fill_target_chars = 2700
            max_one_page_chars = 12000
            if len(resume) < fill_min_chars:
                import copy
                resume_before_fill = resume
                fb_before_fill = copy.deepcopy(fb) if isinstance(fb, dict) else fb
                fb_fill = dict(fb or {})
                fb_fill['allow_repopulate_from_master'] = True
                fb_fill['repopulate_target_chars'] = int(fill_target_chars)
                fb_fill['lock_summary'] = True
                fb_fill.setdefault('preserve_style', True)
                fb_fill['prefer_shorten'] = False
                try:
                    raw_fill = _safe_call(rewrite, jd, normalized_master, resume, fb_fill, timeout=60)
                except Exception:
                    raw_fill = rewrite(jd, normalized_master, resume, fb_fill)
                parsed_fill, remaining_fill = _extract_json_blob_and_rest(raw_fill)
                if parsed_fill is not None:
                    if isinstance(parsed_fill, dict):
                        fb.update(parsed_fill)
                    resume = remaining_fill.strip()
                else:
                    resume = raw_fill
                # Re-evaluate once after fill so the UI score/feedback matches the final text
                try:
                    fb = _safe_call(evaluate, jd, resume, None, timeout=40)
                except Exception:
                    try:
                        fb = evaluate(jd, resume)
                    except Exception:
                        pass
                # Guardrail: rollback only if something went wildly wrong and output exploded.
                if len(resume) > max_one_page_chars:
                    resume = resume_before_fill
                    fb = fb_before_fill
        except Exception:
            pass

        # Final length guardrail (character-based): enforce a hard cap for consistent one-page behavior.
        # Uses section-aware truncation to preserve the structured format.
        try:
            final_char_cap = 2750
            if len(resume) > final_char_cap:
                # Best-effort: ask the trimmer to shorten wording instead of dropping whole bullets.
                try:
                    fb_trim = dict(fb or {})
                    fb_trim['prefer_shorten'] = True
                    fb_trim['trim_to_chars'] = int(final_char_cap)
                    fb_trim['lock_summary'] = True
                    raw_trim = _safe_call(trimmer, jd, normalized_master, resume, fb_trim, timeout=60)
                    parsed_trim, remaining_trim = _extract_json_blob_and_rest(raw_trim)
                    if parsed_trim is not None and isinstance(parsed_trim, dict):
                        fb.update(parsed_trim)
                        resume = remaining_trim.strip()
                    else:
                        resume = raw_trim
                except Exception:
                    pass

                # Deterministic fallback: if we still exceed the cap, truncate section-aware.
                if len(resume) > final_char_cap:
                    resume = _truncate_structured_resume(resume, cap=final_char_cap)
        except Exception:
            pass

        # Emergency-only guardrail: extremely large outputs are capped deterministically.
        try:
            emergency_cap_chars = 12000
            if len(resume) > emergency_cap_chars:
                resume = _truncate_structured_resume(resume, cap=emergency_cap_chars)
        except Exception:
            pass

        st.session_state.result = (resume, fb)
        try:
            _save_persisted_state()
        except Exception:
            pass
        # clear running state and progress after completion
        st.session_state['running'] = False
        try:
            main_progress_placeholder.empty()
        except Exception:
            pass
        st.session_state.pop('progress_pct', None)
        st.session_state.pop('progress_msg', None)
        st.session_state.pop('run_start_ts', None)
        _update_progress(100, "Optimization completed")
        st.session_state['running'] = False
        time.sleep(0.3)
        try:
            main_progress_placeholder.empty()
        except Exception as e:
            _report_error("Failed to clear progress UI", e)

# --- Render results in the main column and recommendations/downloads in the right column ---
with sidebar_col:
    # Feedback section (shows recruiter feedback + bullet scores/change log)
    st.header("Feedback")
    ps = st.session_state.get("prescreen_result") or None
    if ps:
        with st.expander("Pre-screening", expanded=True):
            st.write(f"Skill match: {ps.get('skill_match_pct', 0)}%")
            reasons = ps.get("reasons") or []
            for r in reasons[:3]:
                st.write("- " + r)
            tips = ps.get("tips") or []
            if tips:
                st.markdown("**Tips to improve match**")
                for t in tips[:2]:
                    st.write("- " + t)
    if "result" in st.session_state:
        resume, fb = st.session_state.result
        loop_scores = st.session_state.get("loop_scores") or []
        if loop_scores:
            with st.expander("Loop Scores", expanded=True):
                for entry in loop_scores:
                    st.write(f"iter {entry.get('iter')}: {entry.get('score')}")
        with st.expander("Recruiter Feedback", expanded=True):
            st.write(fb["verbal_feedback"])
            st.markdown("**Strengths**")
            for s in fb["strengths"]:
                st.write("- " + s)
            st.markdown("**Weaknesses**")
            for w in fb["weaknesses"]:
                st.write("- " + w)
            # Show ATS breakdown if available (friendly labels)
            try:
                if fb.get("ats_breakdown"):
                    st.markdown("**ATS score breakdown**")
                    bd = fb.get("ats_breakdown") or {}
                    labels = {
                        "heuristic_score": "Overall heuristic score",
                        "keyword_coverage": "Keyword coverage",
                        "skill_coverage": "Skills match",
                        "bullet_relevance_avg": "Avg. bullet relevance",
                        "quality_score": "Quality (readability & metrics)"
                    }
                    # Prefer a stable ordering
                    for key in ["heuristic_score", "keyword_coverage", "skill_coverage", "bullet_relevance_avg", "quality_score"]:
                        if key in bd:
                            val = bd.get(key)
                            try:
                                sval = f"{float(val):.0f}%"
                            except Exception:
                                sval = str(val)
                            st.write(f"- **{labels.get(key,key)}**: {sval}")
            except Exception:
                pass

        # Change Log UI removed per user request
        # preview/proposed-changes flow removed to simplify workflow and reduce LLM calls

        # Show per-bullet relevance scores (score, reason, action) and suggested skills (always present under Feedback)
        id_action = {}
        id_to_bullet = {}
        cl = fb.get("change_log") if fb else None
        if cl:
            if isinstance(cl, dict) and "change_log" in cl:
                cl_entries = cl.get("change_log") or []
            elif isinstance(cl, list):
                cl_entries = cl
            else:
                cl_entries = []

            for e in cl_entries:
                try:
                    action = e.get("action")
                    ids = e.get("ids") or []
                    if isinstance(ids, list):
                        for _id in ids:
                            id_action[_id] = action
                except Exception:
                    continue

        if fb and fb.get("bullet_scores"):
            for b in fb.get("bullet_scores"):
                bid = b.get("id")
                id_to_bullet[bid] = b.get("bullet")

        with st.expander("Bullet Relevance Scores", expanded=False):
            rows = []
            if fb and fb.get("bullet_scores"):
                for b in fb.get("bullet_scores"):
                    bid = b.get("id")
                    rows.append({
                        "score": b.get("score"),
                        "reason": b.get("reason"),
                        "action": id_action.get(bid, "kept")
                    })
            # If generator provided preliminary bullet_scores, show them under the table as note
            gen_scores = st.session_state.get('generated_bullet_scores', [])
            if gen_scores:
                st.markdown("_Note: preliminary per-bullet scores were produced by the generator; final evaluation may differ._")

            # Render as a clean HTML table (no index column) for consistent UI
            try:
                import html
                def _render_rows_as_table(rows):
                    if not rows:
                        return "<p>No bullet scores available.</p>"
                    headers = ["Score", "Reason", "Action"]
                    parts = ['<table style="width:100%; border-collapse: collapse;">']
                    # header
                    parts.append('<thead><tr>')
                    for h in headers:
                        parts.append(f'<th style="text-align:left; padding:6px; border-bottom:1px solid #ddd;">{html.escape(h)}</th>')
                    parts.append('</tr></thead>')
                    # body
                    parts.append('<tbody>')
                    for r in rows:
                        parts.append('<tr>')
                        parts.append(f'<td style="padding:6px; border-bottom:1px solid #f3f3f3; width:60px">{html.escape(str(r.get("score", "")))}</td>')
                        parts.append(f'<td style="padding:6px; border-bottom:1px solid #f3f3f3;">{html.escape(str(r.get("reason", "")))}</td>')
                        parts.append(f'<td style="padding:6px; border-bottom:1px solid #f3f3f3; width:90px">{html.escape(str(r.get("action", "")))}</td>')
                        parts.append('</tr>')
                    parts.append('</tbody>')
                    parts.append('</table>')
                    return ''.join(parts)

                html_table = _render_rows_as_table(rows)
                st.markdown(html_table, unsafe_allow_html=True)
            except Exception:
                for r in rows:
                    st.write(f"{r['score']}: {r['reason']} (action: {r['action']})")

            # Show deleted bullets (if any) with snippets
            try:
                removed = []
                if cl:
                    for e in cl_entries:
                        try:
                            if e.get("action") == "removed":
                                # prefer bullet_snippet, fallback to id->bullet
                                snippet = e.get("bullet_snippet") or ""
                                if not snippet:
                                    ids = e.get("ids") or []
                                    if ids:
                                        snippet = id_to_bullet.get(ids[0], "")
                                if snippet:
                                    removed.append(snippet)
                        except Exception:
                            continue
                if removed:
                    st.markdown("**Removed bullets:**")
                    for s in removed:
                        st.markdown(f"- {s}")
            except Exception:
                pass
            # Suggested skills produced by the generator (LLM) — always show if present
            try:
                suggested = st.session_state.get('suggested_skills', [])
                if suggested:
                    st.markdown("**Suggested skills (from generator)**")
                    st.markdown(", ".join(suggested))
            except Exception:
                pass

        # (Sidebar download removed — Download DOCX is available under the Optimized Resume preview)

with main_col:
    if "result" in st.session_state:
        resume, fb = st.session_state.result

        st.metric("ATS Score", fb["ats_score"])
        # Show quality_score (human-readability heuristic) if available
        try:
            qv = None
            if fb.get("ats_breakdown") and fb["ats_breakdown"].get("quality_score") is not None:
                qv = fb["ats_breakdown"].get("quality_score")
            if qv is not None:
                st.metric("Quality Score", qv)
        except Exception:
            pass

        with st.expander("Optimized Resume (Preview)", expanded=True):
            # Render preview as Markdown with bolded bullet starts and bolded skill categories
            def _resume_to_markdown(resume_text):
                import re
                def _convert_bold_markers(s: str) -> str:
                    if not s:
                        return s
                    # Replace any {BOLD_START} ... {BOLD_END} with Markdown bold **...**
                    return re.sub(r'\{BOLD_START\}\s*(.*?)\s*\{BOLD_END\}', lambda m: f"**{m.group(1)}**", s, flags=re.DOTALL)

                secs = parse_resume(resume_text)
                parts = []
                # Header
                hdr = secs.get('HEADER', [])
                if hdr:
                    # Avoid duplicating a line that actually belongs to another section
                    # (common when a skills line like "Languages / Tools: ..." is mis-parsed
                    # into HEADER). If any other section contains the same raw line, skip it.
                    hdr_raw = hdr[0].strip()
                    other_has_same = False
                    for k, v in secs.items():
                        if k == 'HEADER':
                            continue
                        for ln in v:
                            if hdr_raw == ln.strip():
                                other_has_same = True
                                break
                        if other_has_same:
                            break
                    if not other_has_same:
                        parts.append('**' + _convert_bold_markers(hdr[0]) + '**')
                        if len(hdr) > 1:
                            parts.append(_convert_bold_markers(hdr[1]))
                        parts.append('')

                for sec in ['PROFESSIONAL SUMMARY', 'SUMMARY', 'EXPERIENCE', 'SKILLS', 'EDUCATION']:
                    if sec in secs:
                        parts.append('**' + sec + '**')
                        # Render EXPERIENCE using structured roles so headers and bullets stay separate
                        if sec == 'EXPERIENCE':
                            # Insert a blank line to keep header separate from the first role
                            parts.append('')
                            try:
                                roles = split_experience(secs.get('EXPERIENCE') or secs.get('PROFESSIONAL EXPERIENCE') or [])
                                for role in roles:
                                    hdr_line = []
                                    if role.get('company'):
                                        hdr_line.append(f"**{role.get('company')}**")
                                    if role.get('title'):
                                        hdr_line.append(role.get('title'))
                                    if role.get('meta'):
                                        hdr_line.append(role.get('meta'))
                                        # Ensure a blank line before the role header so Markdown
                                    # doesn't treat it as a continuation of the previous list item.
                                    if parts and parts[-1].startswith('-'):
                                        parts.append('')
                                    role_hdr = ' | '.join(hdr_line)
                                    parts.append(_convert_bold_markers(role_hdr))
                                    for b in role.get('bullets', []):
                                        parts.append('- ' + _convert_bold_markers(b))
                            except Exception:
                                # fallback to naive line rendering if split_experience fails
                                for line in secs[sec]:
                                    if line.startswith('-'):
                                        parts.append('- ' + _convert_bold_markers(line[1:].strip()))
                                    else:
                                        # If a non-bullet line looks like a pipe-style header and the
                                        # previous rendered line is a bullet, insert a blank line
                                        # so Markdown doesn't attach it to the previous list item.
                                        if '|' in line and parts and parts[-1].startswith('-'):
                                            parts.append('')
                                        parts.append(_convert_bold_markers(line))
                        else:
                            for line in secs[sec]:
                                # bullets
                                if line.startswith('-'):
                                    txt = line[1:].strip()
                                    parts.append('- ' + _convert_bold_markers(txt))
                                else:
                                    # skill categories: bold text before ':'
                                    # If the previous line is a bullet and this line is a non-bullet
                                    # (e.g., a stray role header), insert a blank line first so
                                    # Markdown does not attach it to the prior list item.
                                    # If previous line is a bullet and this is a non-bullet,
                                    # make sure the non-bullet starts on its own line so it
                                    # isn't attached to the previous list item.
                                    if parts and parts[-1].startswith('-') and not line.startswith('-'):
                                        parts.append('')
                                    # Also, if this line contains a pipe-style header (e.g.,
                                    # "Company | Title | Dates"), ensure it is separated from
                                    # any preceding non-blank line to avoid attachment.
                                    if '|' in line and parts and parts[-1].strip() != '':
                                        parts.append('')
                                    if sec == 'SKILLS' and ':' in line:
                                        import re as _re
                                        # Strip HTML tags that may be present in pasted content
                                        clean_line = _re.sub(r'<[^>]*>', '', line)
                                        # Find all Category: value pairs on the same line.
                                        # This captures a label up to the next label or end of line.
                                        # First try to detect known category keywords (robust to missing delimiters)
                                        label_keywords = [
                                            'Languages', 'Language', 'Backend', 'Data', 'Infrastructure', 'AI', 'Engineering',
                                            'Tools', 'Frameworks', 'Databases', 'Cloud', 'DevOps', 'Testing', 'Security',
                                            'Projects', 'Certifications', 'Achievements', 'Practices'
                                        ]
                                        kw_pattern = _re.compile(r'(?i)(' + '|'.join(_re.escape(k) for k in label_keywords) + r')')
                                        kw_matches = list(kw_pattern.finditer(clean_line))
                                        labels = []
                                        if kw_matches:
                                            # If multiple keyword matches map to the same colon, prefer the
                                            # match with the latest start (e.g., prefer 'Tools' over 'Languages'
                                            # in "Languages / Tools: ...").
                                            colon_groups = {}
                                            for km in kw_matches:
                                                colon_idx = clean_line.find(':', km.start())
                                                if colon_idx != -1:
                                                    colon_groups.setdefault(colon_idx, []).append(km.start())
                                            for colon_idx, starts in colon_groups.items():
                                                if len(starts) > 1:
                                                    min_start = min(starts)
                                                    max_start = max(starts)
                                                    mid_segment = clean_line[min_start:max_start]
                                                    # If the keywords are separated by a '/', assume they are a combined
                                                    # category like "Languages / Tools" and include both in the label.
                                                    if '/' in mid_segment:
                                                        start_label = min_start
                                                    else:
                                                        # otherwise prefer the later start (more specific label)
                                                        start_label = max_start
                                                else:
                                                    start_label = starts[0]
                                                cat = clean_line[start_label:colon_idx].strip()
                                                labels.append((start_label, colon_idx, cat))
                                        # If keyword-based detection didn't find labels, fallback to generic colon splits
                                        if not labels:
                                            labels = []
                                            for m in _re.finditer(r'([A-Za-z0-9&/\-\+\s]+?):', clean_line):
                                                labels.append((m.start(1), m.end(1), m.group(1).strip()))
                                        # If we have labels, slice values between them
                                        if labels:
                                            labels = sorted(labels, key=lambda x: x[0])
                                            for i, (_, colon_idx, cat) in enumerate(labels):
                                                start = colon_idx + 1
                                                end = labels[i+1][0] if i + 1 < len(labels) else len(clean_line)
                                                rest = clean_line[start:end].strip().rstrip(',')
                                                rest = rest.strip(' ,')
                                                cats = [c.strip() for c in _re.split(r'/\s*', cat)]
                                                for c in cats:
                                                    if rest:
                                                        parts.append(f"**{_convert_bold_markers(c)}:** {_convert_bold_markers(rest)}")
                                                    else:
                                                        parts.append(f"**{_convert_bold_markers(c)}:**")
                                        else:
                                            # Fallback to previous behavior
                                            cat, rest = line.split(':', 1)
                                            cats = [c.strip() for c in _re.split(r'/\s*', cat)]
                                            for c in cats:
                                                parts.append(f"**{_convert_bold_markers(c)}:** {_convert_bold_markers(rest.strip())}")
                                    else:
                                        parts.append(_convert_bold_markers(line))
                        parts.append('')
                return '\n'.join(parts)
            try:
                md = _resume_to_markdown(resume)
                st.markdown(md, unsafe_allow_html=False)
            except Exception:
                st.code(resume, language="text")

            # Allow downloading the currently shown resume (including persisted results)
            try:
                preview_docx = _get_docx_bytes(resume)
            except Exception:
                preview_docx = None
            if preview_docx:
                st.download_button("Download DOCX", data=preview_docx, file_name="resume.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="download_docx_preview")
            else:
                st.download_button("Download TXT", data=resume.encode('utf-8'), file_name="resume.txt", mime="text/plain", key="download_txt_preview")

            # Recruiter feedback moved to the right column for better alignment
            # (rendered in `sidebar_col` below)

            # Company Career Pages moved out of the Preview (rendered separately below)


        sections = parse_resume(resume)
        roles = split_experience(sections.get('EXPERIENCE') or sections.get('PROFESSIONAL EXPERIENCE') or [])
        merged_skills = merge_skills_a1(sections.get("SKILLS", []))
        # Use SKILLS produced by the optimized resume LLM as-is (do not override categories)
        consolidated_skills = merged_skills

        

        # Prepare files for download and show download buttons in the right column
        with tempfile.TemporaryDirectory() as td:
            docx_path = os.path.join(td, "resume.docx")

            # Use robust helper to produce DOCX bytes (falls back to raw-text DOCX to preserve content)
            preview_docx = _get_docx_bytes(resume)

            # Helper: determine number of pages. Prefer PDF page count if available
            def _estimate_pages(resume_text: str):
                # Estimate pages by word count (approx 500 words per Word page)
                words = len(resume_text.split())
                per_page = 500
                import math
                return math.ceil(words / per_page)

            # If the draft is longer than 1 page, iteratively remove lowest-score bullets
            # using force_remove_ids until it fits or no candidates remain.
            def _ensure_one_page(resume_text: str, sections_obj: dict, roles_obj: list, fb_obj: dict):
                # One-page trimming disabled — return original text and computed pages with no changes.
                pages = _estimate_pages(resume_text)
                return resume_text, pages, [], []

            # Run trimming if needed
            # Read initial resume text to estimate pages (we have 'resume' variable earlier)
            current_resume_text = resume
            # If the resume is too long by character count, remove low-importance bullets
            def _ensure_char_limit(resume_text: str, fb_obj: dict, target_chars: int = 1950, min_chars: int = 1800):
                # Character-limit trimming disabled — preserve resume text as-is and report no removals.
                return resume_text, []

            current_pages = _estimate_pages(current_resume_text)
            # Character-length trimming first (target 1950 chars)
            try:
                if not DISABLE_PROGRAMMATIC_TRIMS and len(current_resume_text) > 1950:
                    current_resume_text, char_removed = _ensure_char_limit(current_resume_text, fb, target_chars=1950, min_chars=1800)
                    if char_removed:
                        # re-evaluate after char removals
                        try:
                            fb = evaluate(jd, current_resume_text)
                        except Exception:
                            pass
                        try:
                            sections = parse_resume(current_resume_text)
                            roles = split_experience(sections.get("EXPERIENCE", []))
                            merged_skills = merge_skills_a1(sections.get("SKILLS", []))
                        except Exception:
                            pass
            except Exception:
                pass
            if not DISABLE_PROGRAMMATIC_TRIMS and current_pages > 1:
                # Ensure fb includes bullet_scores (already set earlier in flow)
                current_resume_text, final_pages, removed, restored = _ensure_one_page(current_resume_text, sections, roles, fb)
                # If changes were made, update resume and recompute sections/merged_skills
                if removed or restored:
                    # re-evaluate and re-score after trimming
                    try:
                        fb = evaluate(jd, current_resume_text)
                    except Exception:
                        pass
                    try:
                        sections = parse_resume(current_resume_text)
                        roles = split_experience(sections.get("EXPERIENCE", []))
                        merged_skills = merge_skills_a1(sections.get("SKILLS", []))
                    except Exception:
                        pass
                    # annotate restored ids in change_log for visibility
                    if restored:
                        extra = []
                        for rid in restored:
                            extra.append({
                                "action": "restored",
                                "ids": [rid],
                                "role": "",
                                "bullet_snippet": "",
                                "reason": "Restored to meet minimum 1-page length"
                            })
                        if fb.get("change_log"):
                            # merge
                            if isinstance(fb["change_log"], list):
                                fb["change_log"].extend(extra)
                            elif isinstance(fb["change_log"], dict) and "change_log" in fb["change_log"]:
                                fb["change_log"]["change_log"].extend(extra)
                        else:
                            fb["change_log"] = extra

            # CHAR-LENGTH TRIMMING: enforce character thresholds using bullet scores
            def _ensure_char_limits(resume_text: str, fb_obj: dict, min_chars: int = 1800, start_threshold: int = 1900, target_max: int = 1965, score_threshold: int = 15, max_removals_per_run: int = 2):
                # Character-length trimming disabled — return original text and no removals/entries.
                return resume_text, [], []

            try:
                # run char trimming if current text exceeds start threshold
                cur_len = len(current_resume_text)
                if cur_len > 1900:
                    # adapt trimming aggressiveness based on current ATS score: be more conservative when ATS is low
                    score_threshold_dynamic = 15
                    max_removals_dynamic = 2
                    try:
                        if fb and fb.get('ats_score') is not None:
                            ats_val = int(fb.get('ats_score'))
                            # if ATS is below the user target, restrict removals and increase score threshold
                            if ats_val < target:
                                score_threshold_dynamic = 40
                                max_removals_dynamic = 0
                                # trimming constrained because ATS is below target
                    except Exception:
                        pass
                    if not DISABLE_PROGRAMMATIC_TRIMS:
                        new_text, removed_by_chars, trim_entries = _ensure_char_limits(current_resume_text, fb, min_chars=1800, start_threshold=1900, target_max=1965, score_threshold=score_threshold_dynamic, max_removals_per_run=max_removals_dynamic)
                    else:
                        new_text, removed_by_chars, trim_entries = current_resume_text, 0, []
                    if removed_by_chars or trim_entries:
                        current_resume_text = new_text
                        try:
                            fb = evaluate(jd, current_resume_text)
                        except Exception:
                            pass
                        try:
                            sections = parse_resume(current_resume_text)
                            roles = split_experience(sections.get("EXPERIENCE", []))
                            merged_skills = merge_skills_a1(sections.get("SKILLS", []))
                        except Exception:
                            pass
                        # annotate removals in change_log
                        extra = []
                        # translate removed_by_chars entries and trim_entries into change_log entries
                        for rid in removed_by_chars:
                            extra.append({
                                "action": "removed",
                                "ids": [rid],
                                "role": "",
                                "bullet_snippet": "",
                                "reason": "Removed to meet one-page target"
                            })
                        for te in trim_entries:
                            extra.append({
                                "action": te.get("action", "merged"),
                                "ids": te.get("ids", []),
                                "role": "",
                                "bullet_snippet": "",
                                "reason": f"Trim action: {te.get('action')} (before {te.get('before_len')} -> after {te.get('after_len')})"
                            })
                        if extra:
                            if fb.get("change_log"):
                                if isinstance(fb["change_log"], list):
                                    fb["change_log"].extend(extra)
                                elif isinstance(fb["change_log"], dict) and "change_log" in fb["change_log"]:
                                    fb["change_log"]["change_log"].extend(extra)
                            else:
                                fb["change_log"] = extra
                        # store trim log for UI if requested
                        if show_trim_logs:
                            st.session_state['trim_log'] = extra
            
            except Exception:
                pass

            # If trimming produced a resume that's too short or left very few bullets,
            # attempt up to two repopulation retries where the model is allowed to
            # re-introduce factual bullets from MASTER (not invent new facts) to
            # restore usefulness for the JD.
            try:
                min_chars = 1800
                # compute current bullets count
                try:
                    total_bullets = sum(len(r.get('bullets', [])) for r in roles)
                except Exception:
                    total_bullets = 0

                def _extract_json_blob_local2(s: str):
                    import json
                    for start_ch in ('{', '['):
                        start_idx = s.find(start_ch)
                        if start_idx == -1:
                            continue
                        stack = []
                        pairs = {'{': '}', '[': ']'}
                        open_ch = start_ch
                        close_ch = pairs[open_ch]
                        for k in range(start_idx, len(s)):
                            if s[k] == open_ch:
                                stack.append(open_ch)
                            elif s[k] == close_ch:
                                stack.pop()
                                if not stack:
                                    candidate = s[start_idx:k+1]
                                    rest = s[k+1:]
                                    try:
                                        parsed = json.loads(candidate)
                                        return parsed, rest
                                    except Exception:
                                        break
                    return None, s

                cur_len = len(current_resume_text)
                repop_attempts = 0
                repop_removed = []
                repop_entries = []
                # allow up to 2 repop attempts; accept partial repopulations that
                # increase content toward min_chars by a modest delta (e.g., +50 chars)
                repop_delta_accept = 50
                try:
                    existing_bullets = total_bullets
                except Exception:
                    existing_bullets = 0

                while repop_attempts < 2 and (cur_len < min_chars or total_bullets <= 1):
                    repop_attempts += 1
                    fb_repop = dict(fb)
                    fb_repop['allow_repopulate_from_master'] = True
                    fb_repop['repopulate_target_chars'] = min_chars
                    # If ATS is substantially below target or resume is very short,
                    # allow conservative factual inference so the model may suggest
                    # plausible bullets (the model is required to mark them as inferred).
                    try:
                        ats_val = fb.get('ats_score') if fb and isinstance(fb, dict) else None
                        if (ats_val is not None and isinstance(ats_val, (int, float)) and int(ats_val) < target) or cur_len < (min_chars - 200):
                            fb_repop['allow_factual_inference'] = True
                            fb_repop['inference_strength'] = 'conservative'
                            # prefer preview so user approves inferred content
                            preview_mode_local = True
                    except Exception:
                        pass
                    try:
                        raw_repop = _safe_call(rewrite, jd, normalized_master, current_resume_text, fb_repop, timeout=60)
                    except Exception:
                        raw_repop = rewrite(jd, normalized_master, current_resume_text, fb_repop)
                    parsed_repop, remaining_repop = _extract_json_blob_local2(raw_repop)
                    tentative = remaining_repop.strip() if parsed_repop is not None else raw_repop

                    # If the model returned structured repopulate_candidates, insert them deterministically
                    if isinstance(parsed_repop, dict) and parsed_repop.get('repopulate_candidates'):
                            candidates = parsed_repop.get('repopulate_candidates')
                            # insert each candidate bullet into the best-matching role in current_resume_text
                            lines = current_resume_text.splitlines()
                            sections_tmp = parse_resume(current_resume_text)
                            exp_lines = sections_tmp.get('EXPERIENCE', [])
                            # naive insertion: for each candidate, find a role whose company string appears in candidate 'role' in change_log or fallback to first role
                            for c in candidates:
                                bullet_text = c.get('bullet') or ''
                                # determine target role
                                target_role_line = None
                                # prefer role names mentioned in change_log if present
                                if parsed_repop.get('change_log') and isinstance(parsed_repop.get('change_log'), list):
                                    for ch in parsed_repop.get('change_log'):
                                        role_label = ch.get('role', '')
                                        if role_label and role_label in '\n'.join(lines):
                                            target_role_line = role_label
                                            break
                                # fallback: try to match company names inside bullet
                                if not target_role_line:
                                    for rl in exp_lines:
                                        # rl is a role header like 'Company | Role | Dates | Full-time'
                                        if rl and rl in '\n'.join(lines):
                                            # choose first EXPERIENCE header as default target
                                            target_role_line = rl
                                            break
                                # if still no target, skip
                                if not target_role_line:
                                    continue
                                # find insertion point: locate the line index of target_role_line in original text
                                try:
                                    idx = None
                                    for i, L in enumerate(lines):
                                        if L.strip() == target_role_line.strip():
                                            idx = i
                                            break
                                    if idx is None:
                                        continue
                                    # after header, find the last bullet line belonging to this role
                                    insert_at = idx + 1
                                    j = idx + 1
                                    while j < len(lines) and (lines[j].startswith('-') or not lines[j].isupper()):
                                        # if next section header encountered, stop
                                        if lines[j].isupper() and lines[j].strip() in sections_tmp:
                                            break
                                        insert_at = j + 1
                                        j += 1
                                    # insert bullet (ensure it starts with '-')
                                    lines.insert(insert_at, '- ' + bullet_text)
                                except Exception:
                                    continue
                            tentative = '\n'.join(lines)
                            parsed_repop['repopulate_applied'] = True
                    # safety checks
                    try:
                        sec_tmp_r = parse_resume(tentative)
                        roles_tmp_r = split_experience(sec_tmp_r.get('EXPERIENCE', []))
                        empty_role_r = any(len(r.get('bullets', [])) == 0 for r in roles_tmp_r)
                        total_bullets = sum(len(r.get('bullets', [])) for r in roles_tmp_r)
                    except Exception:
                        empty_role_r = False
                        total_bullets = 0

                    if empty_role_r:
                        # reject this repop attempt
                        continue
                    # if tentative is acceptable, accept and re-evaluate
                    # Acceptance criteria: any of
                    #  - reaches min_chars
                    #  - adds at least one bullet
                    #  - increases length by at least repop_delta_accept characters
                    increased_len = len(tentative) - cur_len
                    added_bullets = total_bullets - existing_bullets
                    if len(tentative) >= min_chars or added_bullets > 0 or increased_len >= repop_delta_accept:
                        current_resume_text = tentative
                        try:
                            fb = evaluate(jd, current_resume_text)
                        except Exception:
                            pass
                        try:
                            sections = parse_resume(current_resume_text)
                            roles = split_experience(sections.get('EXPERIENCE', []))
                            merged_skills = merge_skills_a1(sections.get('SKILLS', []))
                        except Exception:
                            pass

                        # annotate repop actions into change_log / trim_log
                        extra_r = []
                        if parsed_repop and isinstance(parsed_repop, dict):
                            # if model returned a change_log, include it
                            if 'change_log' in parsed_repop:
                                clitems = parsed_repop.get('change_log') or []
                                for ci in clitems:
                                    extra_r.append(ci)
                            else:
                                # fallback entry
                                extra_r.append({
                                    'action': 'repopulated',
                                    'ids': [],
                                    'role': '',
                                    'bullet_snippet': '',
                                    'reason': 'Repopulated facts from MASTER to restore content'
                                })
                        else:
                            extra_r.append({
                                'action': 'repopulated',
                                'ids': [],
                                'role': '',
                                'bullet_snippet': '',
                                'reason': 'Repopulated facts from MASTER to restore content'
                            })

                        if extra_r:
                            if fb.get('change_log'):
                                if isinstance(fb['change_log'], list):
                                    fb['change_log'].extend(extra_r)
                                elif isinstance(fb['change_log'], dict) and 'change_log' in fb['change_log']:
                                    fb['change_log']['change_log'].extend(extra_r)
                            else:
                                fb['change_log'] = extra_r

                            if show_trim_logs:
                                st.session_state.setdefault('trim_log', [])
                                st.session_state['trim_log'].extend(extra_r)

                        # recompute lengths and decide whether to stop
                        cur_len = len(current_resume_text)
                        if cur_len >= min_chars and total_bullets > 1:
                            break
                # end repop attempts
            except Exception:
                pass
            

            try:
                # Use bytes produced by _get_docx_bytes (ensures fallback behavior)
                docx_bytes = preview_docx

                # Ensure preview and DOCX content are identical: update session result
                try:
                    resume = current_resume_text
                    st.session_state.result = (resume, fb)
                    try:
                        _save_persisted_state()
                    except Exception:
                        pass
                except Exception:
                    pass

                # DOCX is available from the preview download button above; no duplicate button here.
            except Exception as e:
                st.error("Failed to prepare resume file for download.")
                st.exception(e)
    else:
        st.info("Click the 'Run Agent' button under Master Resume to generate an optimized resume.")

# Company Career Pages: always visible, separate from the Optimized Resume preview
try:
    st.markdown("<div style='margin-top:18px; padding-top:12px; border-top:1px solid #eee;'></div>", unsafe_allow_html=True)
    st.markdown("<div style='font-size:18px; font-weight:700; margin-bottom:8px;'>Company Career Pages</div>", unsafe_allow_html=True)
    st.markdown("<div style='font-size:14px; color:#444; margin-bottom:14px;'>Quick links to popular company career pages — open in a new tab.</div>", unsafe_allow_html=True)

    cats = list(CATEGORIZED.items())
    cols = st.columns(3)
    for idx, (cat_name, comp_list) in enumerate(cats):
        col = cols[idx % 3]
        with col:
            st.markdown(f"<div style='font-size:16px; font-weight:700; margin-bottom:10px;'>{cat_name}</div>", unsafe_allow_html=True)
            display_list = comp_list[:12]
            if not display_list:
                st.markdown("<div style='font-size:13px; color:#666;'>No entries</div>", unsafe_allow_html=True)
            for c in display_list:
                name = c.get('name', '')
                url = c.get('url', '')
                st.markdown(f"<div style='font-size:15px; margin-left:6px; margin-bottom:6px;'>• <a href=\"{url}\" target=\"_blank\">{name}</a></div>", unsafe_allow_html=True)
except Exception:
    pass
