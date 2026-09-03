"""Text extraction helpers for uploaded job descriptions."""

from io import BytesIO
from pathlib import Path


class InputExtractionError(ValueError):
    """Raised when an uploaded text document cannot be read."""


# Backward-compatible name retained for callers/tests created with JD upload.
JDExtractionError = InputExtractionError


def _clean_text(text: str) -> str:
    text = (text or "").replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.rstrip() for line in text.split("\n")]
    cleaned = "\n".join(lines).strip()
    if not cleaned:
        raise InputExtractionError("The uploaded file does not contain readable text.")
    return cleaned


def _extract_txt(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1252"):
        try:
            return _clean_text(data.decode(encoding))
        except UnicodeDecodeError:
            continue
    raise InputExtractionError("The text file encoding is not supported. Save it as UTF-8 and try again.")


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document

        document = Document(BytesIO(data))
    except Exception as exc:
        raise InputExtractionError("The DOCX file could not be opened.") from exc

    blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    return _clean_text("\n".join(blocks))


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(data))
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception as exc:
                raise InputExtractionError("Password-protected PDF files are not supported.") from exc
        text = "\n\n".join((page.extract_text() or "") for page in reader.pages)
    except InputExtractionError:
        raise
    except Exception as exc:
        raise InputExtractionError("The PDF file could not be opened or parsed.") from exc
    return _clean_text(text)


def extract_uploaded_text(filename: str, data: bytes) -> str:
    """Extract text from a supported uploaded PDF, DOCX, or TXT file."""
    if not data:
        raise InputExtractionError("The uploaded file is empty.")

    suffix = Path(filename or "").suffix.lower()
    if suffix == ".txt":
        return _extract_txt(data)
    if suffix == ".docx":
        return _extract_docx(data)
    if suffix == ".pdf":
        return _extract_pdf(data)
    raise InputExtractionError("Unsupported file type. Upload a PDF, DOCX, or TXT file.")


def extract_job_description(filename: str, data: bytes) -> str:
    """Backward-compatible JD-specific wrapper."""
    return extract_uploaded_text(filename, data)
