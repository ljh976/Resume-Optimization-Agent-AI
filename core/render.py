
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
import logging
import re

# NOTE: bolding and NLP-based keyword extraction removed per user request.

def _convert_markdown_bold_to_markers(text: str) -> str:
    """Translate simple Markdown bold (**text**) into our explicit marker format."""
    if not text or "**" not in text:
        return text

    pattern = re.compile(r"\*\*(.+?)\*\*")

    def repl(match):
        return "{BOLD_START}" + match.group(1) + "{BOLD_END}"

    return pattern.sub(repl, text)


def _split_marked_segments(text: str):
    """Split a string by {BOLD_START}/{BOLD_END} markers into (segment, bold_bool) pairs.
    Only honors explicit markers; returns original text as a single non-bold segment if markers are absent.
    """
    text = _convert_markdown_bold_to_markers(text)
    if "{BOLD_START}" in text and "{BOLD_END}" in text:
        parts = []
        s = text
        while True:
            si = s.find('{BOLD_START}')
            if si == -1:
                if s:
                    parts.append((s, False))
                break
            if si > 0:
                parts.append((s[:si], False))
            s = s[si + len('{BOLD_START}') :]
            ei = s.find('{BOLD_END}')
            if ei == -1:
                # malformed: treat remainder as non-bold
                parts.append((s, False))
                break
            parts.append((s[:ei], True))
            s = s[ei + len('{BOLD_END}') :]
        # merge adjacent same-flag segments
        merged = []
        for seg, b in parts:
            if not seg:
                continue
            if merged and merged[-1][1] == b:
                merged[-1] = (merged[-1][0] + seg, b)
            else:
                merged.append((seg, b))
        return merged
    return [(text, False)]




def _add_marked_runs(paragraph, text: str, *, font_name: str = "Cambria", font_size: int = 11, bold: bool = False, italic: bool = False):
    """Add runs to a paragraph, honoring {BOLD_START}/{BOLD_END} markers for selective bolding."""
    if text is None:
        return
    for seg_text, seg_bold in _split_marked_segments(text):
        if not seg_text:
            continue
        run = paragraph.add_run(seg_text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bool(seg_bold or bold)
        if italic:
            run.italic = True


def _apply_bullet_orphan_control(text: str, *, tail_words: int = 3) -> str:
    """Reduce cases where the last single word wraps to its own line in DOCX.

    This is a best-effort layout hint: join the last `tail_words` words using
    non-breaking spaces (U+00A0) so Word keeps them together.

    NOTE: We intentionally apply this only to EXPERIENCE bullets.
    """
    if not text:
        return text
    if tail_words < 2:
        return text

    # Avoid touching lines that already contain explicit newlines.
    if "\n" in text or "\r" in text:
        return text

    # Apply even for moderately short bullets where orphan wraps are common.
    if len(text) < 45:
        return text

    parts = text.rsplit(" ", tail_words - 1)
    if len(parts) != tail_words:
        return text
    # Join the tail words with NBSP; keep earlier content unchanged.
    head = parts[0].rstrip()
    tail = "\u00A0".join(p.strip() for p in parts[1:] if p is not None)
    if not head or not tail:
        return text
    return head + " " + tail



def _tight(p, before=0, after=0):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)


def _add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    new_run.append(rPr)
    new_run.append(OxmlElement('w:t'))
    new_run.find(qn('w:t')).text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def render_docx(path, sections, roles, merged_skills):
    doc = Document()

    # Moderate margin reduction to improve one-page fit while keeping readability.
    try:
        for sec in doc.sections:
            sec.top_margin = Inches(0.65)
            sec.bottom_margin = Inches(0.65)
            sec.left_margin = Inches(0.70)
            sec.right_margin = Inches(0.70)
    except Exception:
        pass

    header = sections.get("HEADER", [])
    name = header[0] if len(header) > 0 else ""
    contact = header[1] if len(header) > 1 else ""

    # --- NAME ---
    if name:
        p = doc.add_paragraph()
        _tight(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_marked_runs(p, name, font_size=20, bold=True)

    # --- CONTACT LINE WITH HYPERLINKS ---
    if contact:
        parts = [p.strip() for p in contact.split("|")]
        p = doc.add_paragraph()
        _tight(p, after=4)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for idx, part in enumerate(parts):
            lower = part.lower()
            if "github.com" in lower:
                _add_hyperlink(p, "GitHub", "https://" + lower.replace("https://", "").replace("http://", ""))
            elif "linkedin.com" in lower:
                _add_hyperlink(p, "LinkedIn", "https://" + lower.replace("https://", "").replace("http://", ""))
            else:
                _add_marked_runs(p, part, font_size=11)

            if idx < len(parts) - 1:
                r = p.add_run(" | ")
                r.font.name = "Cambria"
                r.font.size = Pt(11)

    # --- BODY ---
    # Accept both 'SUMMARY' and 'PROFESSIONAL SUMMARY' keys produced by the parser
    for sec in ["SUMMARY","PROFESSIONAL SUMMARY","EXPERIENCE","SKILLS","EDUCATION"]:
        if sec not in sections:
            continue

        p = doc.add_paragraph(sec)
        # Reduce global section-heading spacing so all headers sit closer to content.
        _tight(p, before=3, after=0)
        r = p.runs[0]
        r.font.name = "Cambria"
        r.font.size = Pt(11)
        r.bold = True
        r.underline = True

        if sec == "EXPERIENCE":
            rendered_any_role = False
            for role_idx, role in enumerate(roles):
                p = doc.add_paragraph()
                # Keep first role directly attached to the header; keep later roles compact.
                _tight(p, before=0 if role_idx == 0 else 3, after=0)
                header_parts = []
                if role.get("company"):
                    header_parts.append("{BOLD_START}" + role.get("company") + "{BOLD_END}")
                if role.get("meta"):
                    header_parts.append(role.get("meta"))
                header_line = " | ".join([part.strip() for part in header_parts if part])
                _add_marked_runs(p, header_line)
                rendered_any_role = True

                for b in role["bullets"]:
                    # Render bullets honoring explicit {BOLD_START}/{BOLD_END} markers only
                    p = doc.add_paragraph()
                    _tight(p)
                    p.paragraph_format.left_indent = Pt(18)
                    # bullet symbol
                    r_sym = p.add_run("• ")
                    r_sym.font.name = "Cambria"
                    r_sym.font.size = Pt(11)
                    # add segments preserving bold flags
                    _add_marked_runs(p, _apply_bullet_orphan_control(b))

            # Fallback: if role parsing fails, render raw EXPERIENCE lines so content is never dropped.
            if not rendered_any_role:
                raw_exp = sections.get("EXPERIENCE", []) or sections.get("PROFESSIONAL EXPERIENCE", []) or []
                for idx, line in enumerate(raw_exp):
                    txt = (line or "").strip()
                    if not txt:
                        continue
                    if txt.startswith("-") or txt.startswith("•"):
                        p = doc.add_paragraph()
                        _tight(p)
                        p.paragraph_format.left_indent = Pt(18)
                        r_sym = p.add_run("• ")
                        r_sym.font.name = "Cambria"
                        r_sym.font.size = Pt(11)
                        bullet_txt = txt[1:].strip()
                        _add_marked_runs(p, _apply_bullet_orphan_control(bullet_txt))
                    else:
                        p = doc.add_paragraph()
                        _tight(p, before=0 if idx == 0 else 4, after=0)
                        _add_marked_runs(p, txt)

        elif sec == "SKILLS":
            for line in merged_skills:
                p = doc.add_paragraph()
                _tight(p)
                if ":" in line:
                    cat, rest = line.split(":", 1)
                    _add_marked_runs(p, cat + ": ", bold=True)
                    _add_marked_runs(p, rest.strip())
                elif "：" in line:
                    cat, rest = line.split("：", 1)
                    _add_marked_runs(p, cat + ": ", bold=True)
                    _add_marked_runs(p, rest.strip())
                else:
                    # Preserve non-standard skill lines rather than dropping or crashing.
                    _add_marked_runs(p, line)

        else:
            for line in sections[sec]:
                p = doc.add_paragraph()
                _tight(p)
                _add_marked_runs(p, line)

    doc.save(path)
